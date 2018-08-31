
'''
Created on 2017-4-6

@author: James
'''
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH  
from openpyxl import Workbook
from openpyxl import load_workbook
from datetime import datetime
import pandas as pd
from docx.oxml.ns import qn
import os

# from docx.enum.style import WD_STYLE_TYPE
# document = Document('申济安.docx')
# styles = document.styles
# paragraph_styles = [ s for s in styles if s.type == WD_STYLE_TYPE.PARAGRAPH]
# for style in styles:
#     print(style.name)
#     font = style.font
#     print(font.name)
# 
# print('mission complete')
# exit()

def rmbupper(value):
    map  = [u"零",u"壹",u"贰",u"叁",u"肆",u"伍",u"陆",u"柒",u"捌",u"玖"]
    unit = [u"分",u"角",u"元",u"拾",u"佰",u"千",u"万",u"拾",u"佰",u"千",u"亿",
            u"拾",u"佰",u"千",u"万",u"拾",u"佰",u"千",u"兆"]

    nums = []   #取出每一位数字，整数用字符方式转换避大数出现误差   
    for i in range(len(unit)-3, -3, -1):
        if value >= 10**i or i < 1:
            nums.append(int(round(value/(10**i),2))%10)

    words = []
    zflag = 0   #标记连续0次数，以删除万字，或适时插入零字
    start = len(nums)-3     
    for i in range(start, -3, -1):   #使i对应实际位数，负数为角分
        if 0 != nums[start-i] or len(words) == 0:
            if zflag:
                words.append(map[0])
                zflag = 0
            words.append(map[nums[start-i]])
            words.append(unit[i+2])
        elif 0 == i or (0 == i%4 and zflag < 3): #控制‘万/元’
            words.append(unit[i+2])
            zflag = 0
        else:
            zflag += 1
            
    if words[-1] != unit[0]:    #结尾非‘分’补整字
        words.append(u"整")
    return ''.join(words)

xlsxfilename = '20170720东方般若西湖1号私募投资基金客户信息表.xlsx'
wb = load_workbook(filename = xlsxfilename)
ws = wb.active
productname = ws.cell(row=1, column=1).value
print(productname)
df_file = pd.read_excel(xlsxfilename, header=1)

df_file = df_file[df_file['项目名称'].notnull()]
print(df_file)
df_file = df_file[df_file['项目名称'].str.contains(productname)]
names = list(set(df_file['客户姓名'].tolist()))
category = list(set(df_file['类别'].tolist()))
category.sort()


for username in names :    
    df = df_file[df_file['客户姓名'] == username]
    print(df)
    money = df['金额'].sum()
    moneywan = money*10000
    moneystr = '{:.2f}'.format(moneywan)
    
    document = Document()
    style = document.styles['Normal']
    style.font.name = u'宋体'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')

    style.font.size = Pt(11)
    style.paragraph_format.line_spacing = Pt(33)  #行间距
    style.paragraph_format.space_before = Pt(0)  #上行间距
    style.paragraph_format.space_after = Pt(0)  #下行间距
    paragraph = document.add_paragraph()  
    paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run('{}认购确认书'.format(productname))
    run.font.bold = True    
    run.font.size = Pt(14)
      
      
    document.add_paragraph() 
     
    paragraph = document.add_paragraph()  
    paragraph.add_run('尊敬的委托人')
    run = paragraph.add_run(' {} '.format(username))
    run.font.underline = True        
    paragraph.add_run('先生/女士/公司：')
      
    paragraph = document.add_paragraph()  
    paragraph.add_run('    浙江般若理财服务中心有限公司作为“{}”的管理人，收到了合同编号为'.format(productname))
    run = paragraph.add_run(' {:.0f} '.format(df.iloc[0]['编号']))
    run.font.underline = True              
    paragraph.add_run('号《{}》等相关资料及认购资金。'.format(productname))
      
    document.add_paragraph('合计人民币：') 
      
    paragraph = document.add_paragraph()  
    paragraph.add_run('    （大写）:')
    run = paragraph.add_run(' {} '.format(rmbupper(moneywan)))
    run.font.underline = True              
    paragraph.add_run('；')
      
    paragraph = document.add_paragraph()  
    paragraph.add_run('    （小写）:')
    run = paragraph.add_run(' {} '.format(moneystr))
    run.font.underline = True              
    paragraph.add_run('。')
      
    paragraph = document.add_paragraph()  
    paragraph.add_run('    其中')
    for cate in category[:-1] :
        paragraph.add_run('{}类份额 '.format(cate))
        df_cate = df[df['类别']==cate]
        if(len(df_cate)==0):
            catemoney = 0
        else:
            catemoney = df_cate.iloc[0]['金额']  
        run = paragraph.add_run(' {:.0f} '.format(catemoney))
        run.font.underline = True   
        paragraph.add_run('万元，')    
    paragraph.add_run('{}类份额 '.format(category[-1]))
    df_cate = df[df['类别']==cate]
    if(len(df_cate)==0):
        catemoney = 0
    else:
        catemoney = df_cate.iloc[0]['金额']       
    run = paragraph.add_run(' {:.0f} '.format(catemoney))
    run.font.underline = True  
    paragraph.add_run('万元。')   
      
    paragraph = document.add_paragraph()  
    paragraph.add_run('    “{}”已于'.format(productname))
    founddate = df.iloc[0]['成立时间']
    run = paragraph.add_run(' {:4.0f}年{:2.0f}月{:2.0f}日 '.format(founddate/10000,founddate/100 % 100,founddate%100)) 
    run.font.underline = True              
    paragraph.add_run('成立。')
      
    paragraph = document.add_paragraph()  
    paragraph.add_run('    委托人')    
    run = paragraph.add_run(' {} '.format(username))
    run.font.underline = True  
    paragraph.add_run('认购资金额度为人民币 ')
    run = paragraph.add_run(' {} '.format(moneystr))
    run.font.underline = True  
    paragraph.add_run('，认购的基金份数为 ')
    run = paragraph.add_run(' {} '.format(moneystr))
    run.font.underline = True  
    paragraph.add_run('份，支付认购费用为0元。特此确认。')
      
    document.add_paragraph('    感谢您对般若财富的信任，祝您身体健康、万事如意！') 
    document.add_paragraph() 
    document.add_paragraph() 
      
    paragraph = document.add_paragraph('浙江般若理财服务中心有限公司（签章）')  
    paragraph.paragraph_format.right_indent = Inches(0.3)  
    paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
      
    today = datetime.now();
    todaystr = today.strftime("%Y%m%d")    
    paragraph = document.add_paragraph('{}年{}月{}日'.format(todaystr[0:4],todaystr[4:6],todaystr[6:8]))  
    paragraph.paragraph_format.right_indent = Inches(1.0)  
    paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT    
      
      
    document.save('{}.docx'.format(username))
#     os.system('{}.docx'.format(username))
    
#     break
    

print("mission complete")


